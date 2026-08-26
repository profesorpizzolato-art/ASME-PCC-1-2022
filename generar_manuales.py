import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Aplica color de fondo a una celda de tabla."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def create_styled_doc():
    """Crea un documento docx con la configuración estética institucional."""
    doc = docx.Document()
    
    # Configuración de márgenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Encabezado institucional
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("MENFA Capacitaciones Programa de Integridad de Uniones Bridadas basado en ASME PCC-1-2022")
        hrun.font.name = "Arial"
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(128, 128, 128)
        
    return doc

def add_custom_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102) # Azul Marino
    p.paragraph_format.space_after = Pt(12)

def add_custom_heading1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)

def add_custom_heading2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(51, 102, 153)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

def add_body_p(doc, text, bold_prefix=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = "Arial"
        r_bold.font.size = Pt(11)
        r_bold.font.bold = True
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)

def add_callout(doc, text, title=""):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F0F4F8") # Gris-Azul Claro
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.right_indent = Inches(0.1)
    
    if title:
        rt = p.add_run(title + "\n")
        rt.font.name = "Arial"
        rt.font.size = Pt(11)
        rt.font.bold = True
        rt.font.color.rgb = RGBColor(0, 51, 102)
        
    rtxt = p.add_run(text)
    rtxt.font.name = "Arial"
    rtxt.font.size = Pt(10.5)

# =============================================================================
# MANUAL 1
# =============================================================================
doc1 = create_styled_doc()
add_custom_title(doc1, "MANUAL 1: MARCO REGULATORIO, SELECCIÓN DE COMPONENTES Y LIMPIEZA")

add_custom_heading1(doc1, "1. Contexto Normativo: El Cambio de Paradigma ASME PCC-1-2022")
add_body_p(doc1, "La versión 2022 de la norma ASME PCC-1 (Pressure Boundary Bolted Flange Joint Assembly) introduce la transformación más radical en la historia de la integridad de uniones bridadas. Históricamente, el estándar funcionó como una guía de buenas prácticas orientativa (Guidelines for...). En la edición 2022, la palabra 'Guidelines' se eliminó del título oficial. El texto se reestructuró para establecer procedimientos obligatorios de ensamblaje (Mandatory Assembly Procedures).")

add_callout(doc1, 
    "• SHALL (Debe): Requisito OBLIGATORIO. Su incumplimiento invalida el procedimiento y rechaza el montaje en el informe QA/QC R-2.2-2.\n"
    "• SHOULD (Debería): Recomendación de buena práctica técnica.\n"
    "• MAY (Puede): Permiso o discrecionalidad del especialista armador.",
    "CRITERIO SEMÁNTICO DE CUMPLIMIENTO"
)

add_custom_heading1(doc1, "2. Ecosistema de Normas Interconectadas")
add_body_p(doc1, "ASME B16.5: Define dimensiones de bridas, clases de presión (150 a 2500) y diámetros del círculo de pernos (NPS <= 24).", "• ")
add_body_p(doc1, "ASME B16.47: Extiende la geometría para bridas de gran diámetro (NPS 26 a NPS 60, Series A y B).", "• ")
add_body_p(doc1, "ASME B16.20: Tolerancias dimensionales de juntas espiraladas (SWG), Kammprofile y anillos RTJ.", "• ")
add_body_p(doc1, "ASME Sec. VIII Div. 1 (Apéndice 2): Fórmulas de diseño y factores de junta (m y y).", "• ")
add_body_p(doc1, "EN 1591-4: Estándar de calificación y competencias del personal armador.", "• ")

add_custom_heading1(doc1, "3. Limpieza y Examen Mecánico (Secciones 4 y 5)")
add_custom_heading2(doc1, "Limpieza de Caras (Sección 4)")
add_body_p(doc1, "Remover el 100% de residuos de juntas previas utilizando rascadores de bronce o latón.")
add_body_p(doc1, "Queda estrictamente prohibido el uso de cepillos de acero al carbono en bridas de Acero Inoxidable (SS). Esto evita la contaminación por partículas ferrosas que generan corrosión galvánica y picaduras puntuales (pitting). Utilizar únicamente cepillos con cerdas de acero inoxidable o latón.", "Regla Crítica (Sec. 4.b.2): ")

add_custom_heading2(doc1, "Examen de Recubrimientos en Pernos (Sección 5)")
add_body_p(doc1, "El espesor del recubrimiento (pintura, galvanizado, PTFE) en la superficie de apoyo de las tuercas no debe exceder los 130 µm (5 mils). Un espesor mayor sufre fluencia bajo carga de torque, generando pérdida imprevista de precarga por aplastamiento.")

doc1.save("Manual_1_Marco_Normativo_y_Componentes_Pizzolato.docx")

# =============================================================================
# MANUAL 2
# =============================================================================
doc2 = create_styled_doc()
add_custom_title(doc2, "MANUAL 2: CRITERIOS DE INSPECCIÓN DE CARAS Y EVALUACIÓN DE IMPERFECCIONES")

add_custom_heading1(doc2, "1. Tipos de Caras y Acabados Superficiales")
add_body_p(doc2, "El sello de una unión se logra mediante la deformación controlada de la junta sobre las estriaciones de la cara de la brida (Seating surface).")

add_callout(doc2,
    "• Raised Face (RF) con Junta Espiralada: Ra 3.2 a 6.3 µm (125 - 250 µin)\n"
    "• Raised Face (RF) con Junta Blanda: Ra 3.2 a 12.5 µm (125 - 500 µin)\n"
    "• Ring Type Joint (RTJ) - Ranura Anillo: Ra <= 1.6 µm (63 µin)",
    "RANGOS DE RUGOSIDAD SUPERFICIAL (Ra)"
)

add_custom_heading1(doc2, "2. Matriz de Aceptación y Rechazo de Defectos (Tabla 4.1 / Apéndice C)")
add_body_p(doc2, "ASME PCC-1 Tabla 4.1 clasifica las imperfecciones mecánicas (ralladuras, muescas, picaduras) de acuerdo con su ubicación y profundidad:")

add_custom_heading2(doc2, "Criterios de Evaluación por Geometría del Defecto")
add_body_p(doc2, "Ralladura Radial (Canal de Fuga): Una ralladura que corte en sentido radial el área de sellado genera un canal directo para el fluido. Si la profundidad del defecto d > 0.13 mm (0.005 in) o su extensión radial cruza más del 50% del ancho de contacto de la junta (w), LA CARA SE RECHAZA.")
add_body_p(doc2, "Ralladuras Circunferenciales: Al seguir la dirección del estriado fonográfico, no representan un canal de fuga directo. Profundidad máxima tolerable d <= 0.75 mm (0.030 in).")
add_body_p(doc2, "Picaduras Puntuales (Pitting): Profundidad d <= 0.75 mm y diámetro máximo del alvéolo <= 0.25 * w.")

doc2.save("Manual_2_Inspeccion_Imperfecciones_Caras_Pizzolato.docx")

# =============================================================================
# MANUAL 3
# =============================================================================
doc3 = create_styled_doc()
add_custom_title(doc3, "MANUAL 3: METROLOGÍA DE PERNOS, CÁLCULO DE TARGET TORQUE Y FRICCIÓN")

add_custom_heading1(doc3, "1. Ecuación Fundamental del Torque (ASME PCC-1 Eq. O-3)")
add_body_p(doc3, "En la práctica de ingeniería de campo, la formulación analítica general se simplifica mediante el Factor Nut (K):")

add_callout(doc3, "T = K * d * F_b", "FÓRMULA PRÁCTICA DEL TORQUE")

add_body_p(doc3, "T: Torque objetivo (Target Torque) [N·m o ft·lb]", "• ")
add_body_p(doc3, "K: Factor de fricción global (adimensional)", "• ")
add_body_p(doc3, "d: Diámetro nominal del espárrago [m o in]", "• ")
add_body_p(doc3, "F_b: Precarga objetivo en el perno (Target Bolt Load) [N o lbf]", "• ")

add_custom_heading1(doc3, "2. Sensibilidad del Factor K (Fricción de Lubricante)")
add_body_p(doc3, "El factor K resume la fricción en la rosca y el rozamiento de la cara de la tuerca contra la brida:")
add_body_p(doc3, "Seco / Oxidado (Sin lubricación): K = 0.30 - 0.35 (Baja carga de perno alcanzada).", "• ")
add_body_p(doc3, "Aceite de Motor Ligero: K = 0.20 - 0.22 (Carga media).", "• ")
add_body_p(doc3, "Anti-seize Base Níquel / Cobre (Aprobado): K = 0.14 - 0.16 (Carga Nominal Target).", "• ")
add_body_p(doc3, "Pastas de PTFE / MoS2: K = 0.08 - 0.11 (Sobretensión y riesgo de fluencia del perno).", "• ")

add_body_p(doc3, "Está prohibido aplicar lubricante sobre la junta de sellado. El lubricante deteriora la junta y elimina la fricción necesaria para evitar su eyección.", "Regla Crítica (Sec. 8.b.5): ")

doc3.save("Manual_3_Calculo_Target_Torque_ApendiceO_Pizzolato.docx")

# =============================================================================
# MANUAL 4
# =============================================================================
doc4 = create_styled_doc()
add_custom_title(doc4, "MANUAL 4: PROCEDIMIENTOS DE CAMPO, PATRONES DE AJUSTE Y REGISTRO R-2.2-2")

add_custom_heading1(doc4, "1. Patrones de Ajuste y Secuencia de Apriete (Tabla 3)")
add_body_p(doc4, "Para comprimir la junta de forma paralela y evitar la inclinación de la brida (flange cocking), se aplican las secuencias de torque progresivo:")

add_body_p(doc4, "Pase 1 (Snug-Up): Apriete manual de alineación del 10% al 20% del Target Torque.", "1. ")
add_body_p(doc4, "Pase 2: 30% al 50% del Target Torque en secuencia de Cruz / Estrella.", "2. ")
add_body_p(doc4, "Pase 3: 100% del Target Torque en secuencia de Cruz / Estrella.", "3. ")
add_body_p(doc4, "Pase 4 (Pase Circular): 100% del Target Torque continuo en sentido horario tuerca por tuerca.", "4. ")
add_body_p(doc4, "Pase 5 (Pase Final): 100% del Target Torque circular tras un período de espera de 4 horas (absorbe la relajación por creep inicial).", "5. ")

add_custom_heading1(doc4, "2. Gestión de Calidad QA/QC: La Tabla R-2.2-2 (Joint Assembly Record)")
add_body_p(doc4, "El Apéndice R formaliza las planillas obligatorias para el archivo de integridad de la planta. Cada junta bridada armada en campo debe contar con su correspondiente Formulario Corto R-2.2-2 con firma del Operador y del Inspector QA/QC.")

doc4.save("Manual_4_Procedimientos_Campo_ApendiceR_Pizzolato.docx")

print("✅ ¡Los 4 manuales en Word fueron creados exitosamente en tu carpeta de trabajo!")
