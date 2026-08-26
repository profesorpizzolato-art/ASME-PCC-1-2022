# =============================================================================
# SIMULADOR COMPLETO ASME PCC-1-2022 - ENRUTADOR CENTRAL
# Autoría y Propiedad de la Documentación: Fabricio Pizzolato 
# Institución: MENFA Capacitaciones 
# =============================================================================
import streamlit as st
from datetime import date
import sys
import os
import io
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Configuración global y estética de la plataforma Streamlit (Debe ser la primera instrucción)
st.set_page_config(
    page_title="Simulador ASME PCC-1-2022",
    page_icon="🔧",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Fuerza a Python a buscar e indexar las subcarpetas del directorio actual
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

# Importación de los módulos referenciando la subcarpeta /modules
try:
    from modules.module_0 import render_module_0
    from modules.module_1 import render_module_1
    from modules.module_2 import render_module_2
    from modules.module_3 import render_module_3
    from modules.module_4 import render_module_4
    from modules.module_5 import render_module_5
    from modules.module_6 import render_module_6
except ImportError as e:
    st.error(f"⚠️ Error de Infraestructura: No se pudo cargar un módulo. Detalle: {e}")
    st.stop()


# =============================================================================
# MOTOR GENERADOR DE MANUALES TÉCNICOS EN WORD (.DOCX)
# =============================================================================
def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def create_styled_doc():
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("MENFA Capacitaciones — Programa de Integridad de Uniones Bridadas (ASME PCC-1-2022)")
        hrun.font.name = "Arial"
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(128, 128, 128)
    return doc

def add_custom_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_after = Pt(12)

def add_custom_heading1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)

def add_custom_heading2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(51, 102, 153)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

def add_body_p(doc, text, bold_prefix=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = "Arial"
        r_bold.font.size = Pt(10.5)
        r_bold.font.bold = True
    rtxt = p.add_run(text)
    rtxt.font.name = "Arial"
    rtxt.font.size = Pt(10.5)

def add_callout(doc, text, title=""):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F0F4F8")
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.right_indent = Inches(0.1)
    
    if title:
        rt = p.add_run(title + "\n")
        rt.font.name = "Arial"
        rt.font.size = Pt(10.5)
        rt.font.bold = True
        rt.font.color.rgb = RGBColor(0, 51, 102)
        
    rtxt = p.add_run(text)
    rtxt.font.name = "Arial"
    rtxt.font.size = Pt(10)

def build_manual_bytes(manual_num):
    doc = create_styled_doc()
    
    # Intenta insertar el logo en la portada del manual si existe en el directorio
    if os.path.exists("logo_menfa.png"):
        doc.add_image("logo_menfa.png", width=Inches(1.8))
    
    if manual_num == 1:
        add_custom_title(doc, "MANUAL 1: MARCO REGULATORIO, SELECCIÓN DE COMPONENTES Y LIMPIEZA")
        add_custom_heading1(doc, "1. Contexto Normativo: El Cambio de Paradigma ASME PCC-1-2022")
        add_body_p(doc, "La versión 2022 de la norma ASME PCC-1 (Pressure Boundary Bolted Flange Joint Assembly) introduce procedimientos obligatorios de ensamblaje (Mandatory Assembly Procedures).")
        add_callout(doc, "• SHALL (Debe): Requisito OBLIGATORIO.\n• SHOULD (Debería): Recomendación de buena práctica.\n• MAY (Puede): Permiso o discrecionalidad del especialista armador.", "CRITERIO SEMÁNTICO DE CUMPLIMIENTO")
        add_custom_heading1(doc, "2. Ecosistema de Normas Interconectadas")
        add_body_p(doc, "ASME B16.5: Define dimensiones de bridas, clases de presión (150 a 2500) y diámetros del círculo de pernos (NPS <= 24).", "• ")
        add_body_p(doc, "ASME B16.47: Extiende la geometría para bridas de gran diámetro (NPS 26 a NPS 60).", "• ")
        add_body_p(doc, "ASME B16.20: Tolerancias dimensionales de juntas espiraladas (SWG), Kammprofile y anillos RTJ.", "• ")
        add_custom_heading1(doc, "3. Limpieza de Caras (Sección 4)")
        add_body_p(doc, "Queda estrictamente prohibido el uso de cepillos de acero al carbono en bridas de Acero Inoxidable (SS) para evitar la contaminación ferrosa.", "Regla Crítica (Sec. 4.b.2): ")

    elif manual_num == 2:
        add_custom_title(doc, "MANUAL 2: CRITERIOS DE INSPECCIÓN DE CARAS Y EVALUACIÓN DE IMPERFECCIONES")
        add_custom_heading1(doc, "1. Tipos de Caras y Acabados Superficiales")
        add_callout(doc, "• Raised Face (RF) con Junta Espiralada: Ra 3.2 a 6.3 µm (125 - 250 µin)\n• Raised Face (RF) con Junta Blanda: Ra 3.2 a 12.5 µm (125 - 500 µin)\n• Ring Type Joint (RTJ) - Ranura Anillo: Ra <= 1.6 µm (63 µin)", "RANGOS DE RUGOSIDAD SUPERFICIAL (Ra)")
        add_custom_heading1(doc, "2. Matriz de Aceptación y Rechazo de Defectos (Tabla 4.1)")
        add_body_p(doc, "Ralladura Radial (Canal de Fuga): Si la profundidad del defecto d > 0.13 mm (0.005 in) o su extensión radial cruza más del 50% del ancho de contacto de la junta (w), LA CARA SE RECHAZA.")
        add_body_p(doc, "Ralladuras Circunferenciales: Profundidad máxima tolerable d <= 0.75 mm (0.030 in).")

    elif manual_num == 3:
        add_custom_title(doc, "MANUAL 3: METROLOGÍA DE PERNOS, CÁLCULO DE TARGET TORQUE Y FRICCIÓN")
        add_custom_heading1(doc, "1. Ecuación Fundamental del Torque (ASME PCC-1 Eq. O-3)")
        add_callout(doc, "T = K * d * F_b", "FÓRMULA PRÁCTICA DEL TORQUE")
        add_body_p(doc, "T: Torque objetivo (Target Torque) [N·m o ft·lb]", "• ")
        add_body_p(doc, "K: Factor de fricción global (adimensional)", "• ")
        add_body_p(doc, "d: Diámetro nominal del espárrago [m o in]", "• ")
        add_body_p(doc, "F_b: Precarga objetivo en el perno (Target Bolt Load) [N o lbf]", "• ")
        add_custom_heading1(doc, "2. Sensibilidad del Factor K")
        add_body_p(doc, "Anti-seize Base Níquel / Cobre (Aprobado): K = 0.14 - 0.16 (Carga Nominal Target).", "• ")
        add_body_p(doc, "Está prohibido aplicar lubricante sobre la junta de sellado.", "Regla Crítica (Sec. 8.b.5): ")

    elif manual_num == 4:
        add_custom_title(doc, "MANUAL 4: PROCEDIMIENTOS DE CAMPO, PATRONES DE AJUSTE Y REGISTRO R-2.2-2")
        add_custom_heading1(doc, "1. Patrones de Ajuste y Secuencia de Apriete (Tabla 3)")
        add_body_p(doc, "Pase 1 (Snug-Up): Apriete manual de alineación del 10% al 20% del Target Torque.", "1. ")
        add_body_p(doc, "Pase 2 y 3: 30% a 50% y luego 100% en secuencia de Cruz / Estrella.", "2. ")
        add_body_p(doc, "Pase 4 y 5: 100% del Target Torque continuo en sentido horario tuerca por tuerca (Pase Circular).", "3. ")
        add_custom_heading1(doc, "2. Gestión de Calidad QA/QC: Formulario R-2.2-2")
        add_body_p(doc, "Cada junta bridada armada en campo debe contar con su correspondiente Formulario Corto R-2.2-2 firmado.")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# =============================================================================
# ESTRUCTURA PRINCIPAL DEL ENRUTADOR
# =============================================================================
def main():
    if "user_role" not in st.session_state:
        st.session_state["user_role"] = "Operador"

    # -------------------------------------------------------------------------
    # PANEL LATERAL: Identidad Corporativa, Control de Rol y Navegación
    # -------------------------------------------------------------------------
    # Muestra el logo corporativo MENFA si está disponible en la carpeta local
    if os.path.exists("logo_menfa.png"):
        st.sidebar.image("logo_menfa.png", use_column_width=True)
    else:
        st.sidebar.image("https://img.icons8.com/fluency/96/worker-with-road-cone.png", width=80)

    st.sidebar.title("Simulador ASME PCC-1")
    st.sidebar.markdown("**Autor:** Fabricio Pizzolato")
    st.sidebar.markdown("**Unidad:** IPCL MENFA - UTN")
    st.sidebar.markdown("---")

    # Selector de Perfil de Usuario
    rol_seleccionado = st.sidebar.selectbox(
        "Perfil de Usuario Activo:",
        ["Operador de Campo", "Supervisor QA/QC - Ingeniería"],
        index=0 if st.session_state["user_role"] == "Operador" else 1
    )
    
    st.session_state["user_role"] = "Operador" if "Operador" in rol_seleccionado else "Supervisor"

    st.sidebar.markdown("---")
    
    # Menú de navegación modular
    module_selection = st.sidebar.radio(
        "Seleccione el Módulo de Capacitación:",
        [
            "Módulo 0: Prólogo e Introducción",
            "Módulo 1: Procedimientos de Campo",
            "Módulo 2: Inspección y Defectos",
            "Módulo 3: Base de Datos y Pernos",
            "Módulo 4: Control de Torque y Ajuste",
            "Módulo 5: Ensayos y Desarmado Seguro",
            "Módulo 6: Panel de Evaluación Histórica"
        ],
        index=0
    )
    
    st.sidebar.markdown("---")
    
    # -------------------------------------------------------------------------
    # ZONA RESTRINGIDA: DESCARGA PROTEGIDA DE MANUALES EN WORD (.DOCX)
    # -------------------------------------------------------------------------
    st.sidebar.subheader("🔒 Zona Restringida / Instructor")
    CLAVE_INSTRUCTOR = "Pizzolato2026*"

    password_input = st.sidebar.text_input("Clave de Instructor", type="password")

    if password_input == CLAVE_INSTRUCTOR:
        st.sidebar.success("Acceso concedido — Panel Instructor")
        st.sidebar.markdown("**Descarga de Manuales Word (.docx)**")
        
        st.sidebar.download_button(
            label="📥 Manual 1 (Marco & Limpieza)",
            data=build_manual_bytes(1),
            file_name="Manual_1_Marco_Normativo_Pizzolato.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.sidebar.download_button(
            label="📥 Manual 2 (Inspección Caras)",
            data=build_manual_bytes(2),
            file_name="Manual_2_Inspeccion_Caras_Pizzolato.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.sidebar.download_button(
            label="📥 Manual 3 (Target Torque)",
            data=build_manual_bytes(3),
            file_name="Manual_3_Calculo_Torque_Pizzolato.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.sidebar.download_button(
            label="📥 Manual 4 (Procedimientos Campo)",
            data=build_manual_bytes(4),
            file_name="Manual_4_Procedimientos_Campo_Pizzolato.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    elif password_input != "":
        st.sidebar.error("Clave incorrecta")

    st.sidebar.markdown("---")

    # Referencias Normativas
    st.sidebar.subheader("📖 Referencias Normativas")
    st.sidebar.caption(
        "• **ASME PCC-1-2022:** Pressure Boundary Bolted Flange Joint Assembly\n"
        "• **ASME B16.5 / B16.47:** Steel Pipe Flanges & Fittings\n"
        "• **ASME B16.20:** Metallic Gaskets for Pipe Flanges\n"
        "• **ASME Sec VIII Div 1 App 2:** Rules for Bolted Flange Connections\n"
        "• **EN 1591-4:** Qualification of Personnel Competence"
    )

    st.sidebar.markdown("---")
    st.sidebar.caption(f"Plataforma Educativa v2.5 • {date.today().strftime('%Y')}")

    # -------------------------------------------------------------------------
    # INDICADOR SUPERIOR DE PERFIL ACTIVO
    # -------------------------------------------------------------------------
    col_hdr1, col_hdr2 = st.columns([3, 1])
    with col_hdr1:
        st.caption(f"PERFIL ACTIVO: **{rol_seleccionado.upper()}**")
    with col_hdr2:
        if st.session_state["user_role"] == "Supervisor":
            st.warning("🔒 MODO SUPERVISIÓN / INGENIERÍA HABILITADO")
        else:
            st.success("👷 MODO OPERACIÓN DE CAMPO HABILITADO")

    # -------------------------------------------------------------------------
    # ENRUTAMIENTO LÓGICO: Renderizado Condicional de Módulos
    # -------------------------------------------------------------------------
    if "Módulo 0" in module_selection:
        render_module_0()
        
    elif "Módulo 1" in module_selection:
        render_module_1()
        
    elif "Módulo 2" in module_selection:
        render_module_2()
        
    elif "Módulo 3" in module_selection:
        render_module_3()
        
    elif "Módulo 4" in module_selection:
        render_module_4()
        
    elif "Módulo 5" in module_selection:
        render_module_5()
        
    elif "Módulo 6" in module_selection:
        render_module_6()

if __name__ == "__main__":
    main()
